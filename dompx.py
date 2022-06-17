import io
import re

from re import Pattern
from typing import Iterable
from functools import partial

from docx import Document
from docx.shared import Mm
from docx.text.run import Run
from docx.text.paragraph import Paragraph


def paragraphs(document: Document) -> Paragraph:
    """Paragraph generator for the document

    Parameters
    ----------
    document : Document

    Yields
    ------
    Paragraph

    """
    # first are document level paragaphs
    for paragraph in document.paragraphs:
        yield paragraph

    # we have also paragraph hidden in the document level tables
    yield from table_paragraphs(document.tables)

    # header level paragraphs goes here
    header = document.sections[0].header
    for paragraph in header.paragraphs:
        yield paragraph

    # header level table paragraphs goes here
    yield from table_paragraphs(header.tables)

    # footer level paragraphs goes here
    footer = document.sections[0].footer
    for paragraph in footer.paragraphs:
        yield paragraph

    # footer level table paragraphs goes here
    yield from table_paragraphs(footer.tables)


def table_paragraphs(tables: Iterable) -> Paragraph:
    """Extracting table-level paragraphs which are hidden in the table cells

    Parameters
    ----------
    tables : Iterable

    Yields
    ------
    Paragraph

    """
    for table in tables:
        for col in table.columns:
            for cell in col.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

                # but wait, there's more! what about tables hidden in the
                # table cells?
                yield from table_paragraphs(cell.tables)


def domp(paragraph: Paragraph, doc: Document, token: Pattern, data: dict):
    """Domp. Domp-Domp. Domp...

    Parameters
    ----------
    paragraph : Paragraph
    doc : Document
    token : Pattern
    data : dict

    """
    # do not even continue with the runs if paragraph text doesn't contain
    # token
    if not token.search(paragraph.text):
        return

    # paragraphs consits from the runs and run can contain tokens, which we are
    # going to replace with values
    for item in paragraph.runs:
        for expr, mod in (token.findall(item.text) or []):
            # modifier defines a method of the document manipulation, e.g.
            # replace, img, tbl, etc.. The default fallback is replace.
            globals().get(mod[1:], replace)(doc, item, expr, mod, data)


def replace(doc: Document, run: Run, expr: str, mod: str, data: dict):
    """Embbed text into the document

    Parameters
    ----------
    doc : Document
    run : Run
    expr : str
    mod : str
    data : dict

    """
    # first we need to compile expression against the data to get embbedable
    # value, and then we need to replace expression in the run we the value
    run.text = run.text.replace(f'{expr}{mod}', str(compile_expr(expr, data)))


def img(doc: Document, run: Run, expr: str, mod: str, data: dict):
    """Embbed image into the document

    Parameters
    ----------
    doc : Document
    run : Run
    expr : str
    mod : str
    data : dict

    """
    # in case of images we don't need expression in the run, hence replacing it
    # with the empty string
    run.text = run.text.replace(f'{expr}{mod}', '')
    # image "value" here should be a path of the image
    if picture := compile_expr(expr, data):
        path, width, height = None, None, None

        if isinstance(picture, str):
            path = picture
        elif isinstance(picture, tuple):
            path, width, height = picture
        else:
            return

        run.add_picture(
            path,
            width=(width and Mm(width)),
            height=(height and Mm(height))
        )


def tbl(doc: Document, run: Run, expr: str, mod: str, data: dict):
    """Embbed table into the document

    Parameters
    ----------
    doc : Document
    run : Run
    expr : str
    mod : str
    data : dict

    """
    # in case of table we don't need expression in the run, hence replacing it
    # with the empty string
    run.text = run.text.replace(f'{expr}{mod}', '')
    if matrix := compile_expr(expr, data):
        # as for now we are supporting strict structured matrix like data, e.g.
        # list of lists
        if not isinstance(matrix, list) and not isinstance(matrix[0], list):
            return

        # create a table in document with matrix dimensions
        table = doc.add_table(len(matrix), len(matrix[0]))
        table.style = 'Table Grid'

        # and populate the table with the matrix values
        for ridx, row in enumerate(matrix):
            for cidx, cell in enumerate(row):
                table.cell(ridx, cidx).text = str(cell)

        run.element.addnext(table._tbl)


def compile_expr(expr: str, data: dict) -> str:
    """Compiling expression matched by regexp against data

    Parameters
    ----------
    expr : str
    data : dict

    Returns
    -------
    str

    """
    try:
        # we need to remove @, {, } symbols to work with pure python expression
        val = eval(expr[1:].strip('{}'), {}, data)  # nosec
    except NameError:
        return expr  # in case of absent expression in data do nothing

    return val


def compile(document: bytes, data: dict) -> bytes:
    """Compile document by parsing special variables/expressions for replacing
    them with the provided data

    Parameters
    ----------
    document : bytes
    data : dict

    Returns
    -------
    bytes

    """
    # regular expression for extracting wheezy.template like variables and
    # expressions from the docx document: ex. @{data['key']}!ss
    token = re.compile(r'(@{?[\w\.\[\]\'\"\(\)]+}?)(![a-z]+)?')

    # loading and parsing docx document - lxml under the hood
    doc = Document(document)

    # this is the main loop, it iterates over all paragraphs found in document
    # and calling domp function on them. domp is responsible for tokenizing and
    # replacing expressions with actual data
    any(map(partial(domp, doc=doc, token=token, data=data), paragraphs(doc)))

    # wrapping the result up in a byte stream here and sending it out
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


if __name__ == '__main__':
    data = {
        'first_names_u': 'Garegin',
        'last_names_u': 'Srvandztyan',
        'signatory_person_name': 'Կիրակոս Սարիբեկյան',
        'tracking': 'ERWE-SDFS-213F-2F8F',
        'qr': ('qr.png', 10, None),
        'stamp': ('stamp.png', 20, None),
        'signature': ('sign.jpg', 15, 15)
    }

    with open('demo_qr.docx', 'rb') as indoc, open('out.docx', 'wb') as outdoc:
        outdoc.write(
            compile(indoc, data).read()
        )
